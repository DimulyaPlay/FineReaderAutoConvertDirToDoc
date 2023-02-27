import os
import sys
import subprocess
from time import sleep
import json
import tempfile
import docx
from docx.shared import Cm
import fitz
from datetime import datetime


class main_agregator:
    def __init__(self):
        self.init_application_path()
        self.config_path = os.path.join(self.application_path, 'config.json')
        self.read_create_config()

    def init_application_path(self):
        if getattr(sys, 'frozen', False):
            application_path = os.path.dirname(sys.executable)
        elif __file__:
            application_path = os.path.dirname(__file__)
        self.application_path = application_path
        self.icon = self.resource_path('document-convert.png')

    def resource_path(self, relative_path):
        """ Get absolute path to resource, works for dev and for PyInstaller """
        try:
            # PyInstaller creates a temp folder and stores path in _MEIPASS
            base_path = sys._MEIPASS
        except Exception:
            base_path = os.path.abspath(".")

        return os.path.join(base_path, relative_path)

    def read_create_config(self):
        default_configuration = {
            'finecmd_path': r'C:\Program Files (x86)\ABBYY FineReader 15\FineCmd.exe',
            'input_folder_txt': r'C:\scan',
            'input_folder_img': r'C:\scan',
            'export_wordimages_folder': r'C:\scan\word',
            'export_wordtext_folder': r'C:\scan\word',
            'delay': 1,
            'toText': True,
            'toImg': False,
            'toTextUsePrefix': False,
            'toImgUsePrefix': False,
            'toTextPrefix': '2txt',
            'toImgPrefix': '2img',
            'dpi': 75
        }
        if os.path.exists(self.config_path):
            try:
                with open(self.config_path, 'r') as configfile:
                    self.cfg = json.load(configfile)
            except Exception as e:
                print(e)
                os.remove(self.config_path)
                self.cfg = default_configuration
                with open(self.config_path, 'w') as configfile:
                    json.dump(self.cfg, configfile)
        else:
            self.cfg = default_configuration
            with open(self.config_path, 'w') as configfile:
                json.dump(self.cfg, configfile)
        self.finecmd_path = self.cfg['finecmd_path']
        self.input_path_txt = self.cfg['input_folder_txt']
        self.input_path_img = self.cfg['input_folder_img']
        self.same_input_paths = self.input_path_img == self.input_path_txt
        self.output_path_text = self.cfg['export_wordtext_folder']
        self.output_path_img = self.cfg['export_wordimages_folder']
        self.same_output_paths = self.output_path_text == self.output_path_img
        self.delay = self.cfg['delay']
        self.toText = self.cfg['toText']
        self.toImg = self.cfg['toImg']
        self.toTextUsePrefix = self.cfg['toTextUsePrefix']
        self.toImgUsePrefix = self.cfg['toImgUsePrefix']
        self.toTextPrefix = self.cfg['toTextPrefix']
        self.toImgPrefix = self.cfg['toImgPrefix']
        self.dpi = self.cfg['dpi']

    def write_config_to_file(self):
        with open(self.config_path, 'w') as configfile:
            json.dump(self.cfg, configfile)

    def save_to_docx_as_img(self, filepath, dpi):
        if not filepath.endswith('.pdf'):
            return
        pdffile = fr"{filepath}"
        filename = os.path.basename(filepath)
        pdfdoc = fitz.open(pdffile)
        zoom = 2
        mat = fitz.Matrix(zoom, zoom)
        count = len(pdfdoc)
        docxdoc = docx.Document()
        for section in docxdoc.sections:
            section.top_margin = Cm(0.5)
            section.bottom_margin = Cm(0.5)
            section.left_margin = Cm(0.5)
            section.right_margin = Cm(0.5)
        temp_images = []
        for i in range(count):
            fd, outpath = tempfile.mkstemp('.jpg')
            temp_images.append(outpath)
            os.close(fd)
            page = pdfdoc.load_page(i)
            pix = page.get_pixmap(dpi=dpi)
            pix.save(outpath)
            if page.rect.width > page.rect.height:
                docxdoc.add_picture(outpath, width=Cm(21))
            else:
                docxdoc.add_picture(outpath, height=Cm(27))
        pdfdoc.close()
        [os.unlink(i) for i in temp_images]
        docxdoc.save(fr"{self.output_path_img}\\{filename}_images.docx")

    def save_to_doc_as_text(self, filepath):
        filepath_out = self.output_path_text + r'\\' + os.path.basename(filepath) + '_text.doc'
        subprocess.call(self.finecmd_path + ' "' + filepath + '" /lang russian english /out "' + filepath_out + '" /quit')


def ctime():
    return datetime.today().strftime("%Y-%m-%d %H:%M:%S")

